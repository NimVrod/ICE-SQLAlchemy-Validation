"""Fill ICE_projects_template (1).docx with project achievements."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "ICE_projects_template (1).docx"
OUTPUT = ROOT / "ICE_projects_report_filled.docx"


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def clear_paragraph(paragraph: Paragraph) -> None:
    set_paragraph_text(paragraph, "")


def insert_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    set_paragraph_text(new_para, text)
    return new_para


def find_paragraph(doc: Document, text: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == text:
            return i
    raise ValueError(f"Paragraph not found: {text!r}")


def find_heading(doc: Document, text: str) -> int:
    return find_paragraph(doc, text)


def replace_section(
    doc: Document,
    heading_text: str,
    next_heading_text: str | None,
    lines: list[str],
    body_style: str = "Treść Praca Dyplomowa",
    heading_style: str | None = None,
) -> None:
    """Replace all body paragraphs between two headings."""
    start = find_heading(doc, heading_text)
    end = find_heading(doc, next_heading_text) if next_heading_text else len(doc.paragraphs)

    if heading_style:
        doc.paragraphs[start].style = heading_style

    for idx in range(end - 1, start, -1):
        element = doc.paragraphs[idx]._element
        element.getparent().remove(element)

    current = doc.paragraphs[start]
    for line in lines:
        current = insert_after(current, line, body_style)


def append_section_content(
    doc: Document,
    heading_text: str,
    intro_lines: list[str],
    subsections: list[tuple[str, list[str]]],
    body_style: str = "Treść Praca Dyplomowa",
    sub_style: str = "Heading 2",
) -> None:
    current = doc.paragraphs[find_heading(doc, heading_text)]
    for line in intro_lines:
        current = insert_after(current, line, body_style)
    for sub_heading, lines in subsections:
        current = insert_after(current, sub_heading, sub_style)
        for line in lines:
            current = insert_after(current, line, body_style)


def append_body(
    doc: Document,
    heading_text: str,
    items: list,
    body_style: str = "Treść Praca Dyplomowa",
) -> None:
    """Append body paragraphs after a heading.

    Each item is either a plain string, or a (label, text) tuple where the
    label is rendered as a bold lead-in within the same paragraph.
    """
    anchor = doc.paragraphs[find_heading(doc, heading_text)]
    current_p = anchor._p
    parent = anchor._parent
    for item in items:
        new_p = deepcopy(anchor._p)
        current_p.addnext(new_p)
        current_p = new_p
        para = Paragraph(new_p, parent)
        para.style = body_style
        for run in list(para.runs):
            run._element.getparent().remove(run._element)
        if isinstance(item, tuple):
            label, text = item
            bold_run = para.add_run(label)
            bold_run.bold = True
            para.add_run(" " + text)
        else:
            para.add_run(item)


def fill_table(doc: Document) -> None:
    table = doc.tables[0]
    rows = [
        ("Marker", "Check", "SQL strategy"),
        ("Required()", "Value must not be NULL", "WHERE column IS NULL"),
        ("Unique()", "Column values must be unique", "GROUP BY … HAVING COUNT(*) > 1"),
        ("UniqueComposite", "Multi-column uniqueness", "Duplicate join on column tuple"),
        ("FloatMin / FloatMax", "Numeric bounds", "WHERE value < min OR value > max"),
        ("NoNaN()", "Reject IEEE NaN in floats", "SQL prefilter + Python math.isnan"),
    ]
    while len(table.rows) < len(rows):
        table.add_row()
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            table.rows[r].cells[c].text = value
        for c in range(len(row_data), len(table.rows[r].cells)):
            table.rows[r].cells[c].text = ""


def main() -> None:
    doc = Document(TEMPLATE)

    set_paragraph_text(doc.paragraphs[2], "ICE SQLAlchemy Validation")
    for idx, name in zip(
        range(4, 9),
        [
            "1. Hubert Dec",
            "2. Patryk Dziki",
            "3. Adrian Czapka",
            "4. Dominik Dziedzic",
            "5. Beniamin Sujka",
        ],
    ):
        set_paragraph_text(doc.paragraphs[idx], name)
    set_paragraph_text(doc.paragraphs[9], "Supervisor: [to be completed]")
    set_paragraph_text(doc.paragraphs[10], "Rzeszów, June 2026")

    replace_section(
        doc,
        "Introduction",
        "Section Project description",
        [
            (
                "Database migrations often complete without error while silently corrupting data. "
                "Public postmortems document cases where row counts matched, logs stayed clean, "
                "and tooling reported success — yet timestamps shifted, categories were remapped, "
                "or customer records were associated with the wrong accounts [1–6]. "
                "The ICE SQLAlchemy Validation project addresses a structural part of this problem: "
                "verifying that rows already loaded into a target database still satisfy the business "
                "rules declared on SQLAlchemy ORM models."
            ),
            (
                "The project team consists of five students from the ICE programme: "
                "Hubert Dec, Patryk Dziki, Adrian Czapka, Dominik Dziedzic, and Beniamin Sujka. "
                "Work was organised in two sprints. Sprint I produced a research report on migration "
                "data-quality incidents and failure modes (WriteUps/Error_refrences_SprintI). "
                "Sprint II implemented a Python library in src/ that audits migrated tables using "
                "set-based SQL, pandas reporting, and pytest integration."
            ),
            (
                "The assumed methodology combined literature and incident review, iterative design "
                "of validation markers on ORM columns, test-driven implementation, and demonstration "
                "on synthetic dirty datasets generated with Faker. The solution is an open-source "
                "library rather than a commercial product; it is intended for engineering teams "
                "running bulk loads, ETL pipelines, or schema migrations outside normal ORM insert paths."
            ),
        ],
    )

    replace_section(
        doc,
        "Section Project description",
        "Team members",
        [
            (
                "The project delivers a migration data-quality library for SQLAlchemy 2.x. "
                "Validation rules are declared once on each mapped column using typing.Annotated "
                "markers (Required, Nullable, Unique, UniqueComposite, FloatMin, FloatMax, NoNaN). "
                "The MigrationValidator engine reads those declarations and scans live database "
                "tables, returning every violating row."
            ),
            (
                "Core components in src/: metadata.py defines markers; parser.py resolves "
                "Annotated metadata and builds validator objects; validators/ implements "
                "field-level checks; engine.py orchestrates SQL-first validation in two passes "
                "(field rules, then duplicate rules); pandas_.py converts results to DataFrames "
                "and HTML reports; pytest_validation_plugin.py exposes CI-friendly assertions."
            ),
            (
                "The library targets structural failure modes identified in the Sprint I report: "
                "silent NULLs after bulk load, numeric out-of-range values, duplicate records from "
                "non-idempotent ETL re-runs, composite-key collisions, and NaN artefacts from "
                "pandas/SQLite loads. It does not yet cover semantic drift (timezones, category "
                "remapping, encoding), which remain future work."
            ),
            (
                "Deliverables include the src package, a pytest suite under Tests/, runnable "
                "examples (validate_dirty_database.py, migration_validation_demo.py, Jupyter "
                "notebook), presentation materials in WriteUps/, and project documentation in README.md."
            ),
        ],
    )

    replace_section(
        doc,
        "Team members",
        "Section Business plan",
        [
            (
                "Hubert Dec — project coordination, research synthesis from migration incident "
                "literature, architecture of the validation engine, and integration of pandas/HTML "
                "reporting."
            ),
            (
                "Patryk Dziki — ORM metadata design (Annotated markers), parser implementation, "
                "and declaration tests ensuring single-source-of-truth rules on models."
            ),
            (
                "Adrian Czapka — SQL validator implementations (null, float bounds, duplicates), "
                "streaming query design with yield_per for large tables."
            ),
            (
                "Dominik Dziedzic — pytest plugin, CI entry point, demo scripts, and Faker-based "
                "dirty-database examples used in validation_report.ipynb."
            ),
            (
                "Beniamin Sujka — Sprint I research report and slides (Typst/PDF/PPTX), "
                "competitive analysis, risk documentation, and technical writing for this report."
            ),
        ],
    )

    replace_section(
        doc,
        "Section Business plan",
        "Analysis of competitive solutions/analysis of market (mandatory)",
        [],
    )
    append_body(
        doc,
        "Section Business plan",
        [
            (
                "Although ICE SQLAlchemy Validation is delivered as an academic project, it is "
                "scoped, costed, and positioned as if it were a commercial product entering an "
                "existing data-tooling market. The central business thesis is simple: a database "
                "migration that reports success can still corrupt data silently, and the cost of "
                "discovering that corruption in production is orders of magnitude larger than the "
                "cost of catching it before go-live [1–6]. The product turns an abstract, "
                "hard-to-budget risk into a concrete, repeatable, low-cost quality gate."
            ),
            (
                "Value proposition.",
                "The product enforces data-quality rules that are defined once, alongside the data "
                "model, against the records that already exist in the database after a migration — "
                "not only against new data as it is entered. It targets the silent corruption that "
                "standard checks miss, adds negligible effort to run, keeps all data inside the "
                "customer's own environment, and produces an auditable pass/fail sign-off for "
                "cut-over.",
            ),
            (
                "Cost–benefit analysis (ROI).",
                "The cost of adoption is small: defining the rules on existing models is a matter of "
                "hours, the validation runs in seconds as part of the normal build, and there is no "
                "separate platform to license or operate. The benefit is the avoided cost of a "
                "data-corruption incident, and the Sprint I research shows how large that gap is. "
                "TSB Bank's 2018 migration corrupted around 1.3 billion records, cost approximately "
                "£330 million, drew a £48.65 million regulatory fine, lost about 80,000 customers "
                "and generated 225,492 complaints [4]. A migration between database engines silently "
                "shifted 87% of timestamps and was reported by a customer rather than by monitoring "
                "[3]; an insurance migration mis-dated 127,000 historic policies and was noticed "
                "three weeks after go-live [2]; and an aviation load-sheet error understated aircraft "
                "mass by 1,244 kg [6]. Preventing even one mid-size incident repays the entire "
                "adoption cost many times over: a failed check before go-live costs an engineer about "
                "an hour, against remediation costs measured in millions.",
            ),
            (
                "Revenue and licensing model.",
                "The project follows an open-core model. The core tool is released under a permissive "
                "open-source (MIT) licence to maximise adoption and remove procurement barriers — the "
                "strongest lever for entering a market dominated by established commercial suites. "
                "Future commercial tiers could include paid support and service-level agreements, an "
                "enterprise pack of advanced checks (time-zone and semantic drift, category "
                "remapping, encoding, and cross-table integrity), and a hosted reporting dashboard "
                "with historical baselines. Pricing would be set per team or per project, well below "
                "the cost of a single avoided incident and below incumbent data-observability "
                "platforms.",
            ),
            (
                "Go-to-market and adoption plan.",
                "Adoption is driven first through open-source distribution and runnable "
                "demonstrations, then made durable by embedding the check into a team's existing "
                "release pipeline so every future migration is validated by default. In practice the "
                "tool is run after a bulk data load, before user-acceptance testing, after any fix "
                "that touches legacy data, as a required sign-off before production cut-over, and as "
                "a smoke test immediately afterwards. The guiding principle is to validate at full "
                "production volume rather than on a small sample, because — as in the TSB case — the "
                "edge cases tend to live exactly where smaller test environments never reach.",
            ),
        ],
    )

    replace_section(
        doc,
        "Analysis of competitive solutions/analysis of market (mandatory)",
        "Risk analysis (mandatory)",
        [],
    )
    append_body(
        doc,
        "Analysis of competitive solutions/analysis of market (mandatory)",
        [
            (
                "Market and target customers.",
                "Database migration and modernisation — cloud re-platforming, database-engine swaps, "
                "and billing or ERP consolidation — is a large and recurring area of IT spend, and "
                "Python with SQLAlchemy is one of the most common application stacks involved. The "
                "primary market is therefore engineering teams that perform bulk data loads or "
                "schema migrations outside the application's normal data-entry path. A high-value "
                "secondary market is regulated industries — banking, insurance, healthcare, aviation "
                "— where a silent data error is simultaneously a correctness failure, a potential "
                "privacy breach, and a compliance violation. A third segment is data-platform teams "
                "and consultancies that run migrations for clients and need a defensible, documented "
                "quality gate to sign off a cut-over.",
            ),
            (
                "Competitive landscape.",
                "Several categories of solution already exist, but none closes this specific gap. "
                "Built-in database constraints only protect data when they are enabled and the data "
                "is actually re-checked, yet they are frequently switched off during bulk imports "
                "for performance and not re-applied. Application-level validation only runs on new "
                "data entered through the application, so it never sees records loaded in bulk. "
                "Simple record-count comparisons confirm that data moved, not that each value is "
                "correct. General data-quality and observability platforms (such as Great "
                "Expectations, Soda, dbt tests, or Monte Carlo) are powerful but require a separate "
                "set of rules to be authored and maintained apart from the application, which can "
                "drift out of step with it, and the commercial options carry licensing and operating "
                "cost that is hard to justify for a focused migration sign-off.",
            ),
            (
                "Differentiation and positioning.",
                "ICE SQLAlchemy Validation is intentionally narrow and deep. Its rules live with the "
                "data model itself, giving a single source of truth with nothing to keep in sync; it "
                "checks the records already in the database at full volume; and it runs as part of "
                "the existing release pipeline rather than as a new platform to operate. It does not "
                "attempt to replace a full observability suite, but it decisively wins the specific "
                "job of proving that migrated data satisfies its own rules before cut-over, at a "
                "fraction of the setup and cost. Its open-source core removes the procurement barrier "
                "that typically protects incumbent commercial products. Table 4.1 summarises the "
                "available checks.",
            ),
        ],
    )

    replace_section(
        doc,
        "Risk analysis (mandatory)",
        "Subsubsection if necessary",
        [],
    )
    append_body(
        doc,
        "Risk analysis (mandatory)",
        [
            (
                "The principal risks are set out below, each with its mitigation, covering both "
                "commercial risks and product risks. The most important strategic risk to manage is "
                "false confidence: a clean validation run must never be presented as proof that the "
                "data is fully correct, which is why the product is positioned as one gate that "
                "complements, rather than replaces, a full source-to-target reconciliation."
            ),
            (
                "Adoption risk.",
                "Teams may under-estimate the likelihood of silent corruption until they are "
                "affected by it. Mitigation: lead with concrete, costed incidents and a near-zero "
                "effort first trial, so the value is evident before any commitment.",
            ),
            (
                "Competitive risk.",
                "An established data-quality platform could add similar model-aware validation. "
                "Mitigation: keep the open-source core free and frictionless, move quickly on the "
                "advanced enterprise checks, and own the rules-live-with-the-model experience that "
                "broad platforms are structurally reluctant to match.",
            ),
            (
                "Sustainability risk.",
                "An open-core project depends on ongoing maintenance and community trust. "
                "Mitigation: a deliberately small, well-tested codebase, a documented way to add new "
                "checks, and an optional paid support tier to fund continued maintenance.",
            ),
            (
                "Performance risk.",
                "Validating very large tables could be costly. Mitigation: the database returns only "
                "the records that breach a rule, and results are processed in batches, so the check "
                "remains efficient at production scale.",
            ),
            (
                "Coverage risk.",
                "The current version addresses structural rules only; more complex failures such as "
                "time-zone shifts, category remapping, encoding corruption, and broken links between "
                "tables are not yet implemented. Mitigation: the tool was designed so these "
                "additional checks can be added without re-engineering it, and they head the "
                "enterprise roadmap.",
            ),
            (
                "False-confidence risk.",
                "A clean result does not by itself prove that every migrated value still means what "
                "it did in the source. Mitigation: position the tool explicitly as a quality gate "
                "and recommend a source-to-target comparison for the most critical fields alongside "
                "it.",
            ),
        ],
    )

    replace_section(
        doc,
        "Subsubsection if necessary",
        "Figures and tables (if necessary)",
        [
            (
                "The market, competitive, and risk analysis above is grounded in the documented "
                "migration incidents collected during Sprint I [1–6], together with the recognised "
                "limitations of application-level validation [7] and the project's own "
                "implementation and test evidence [8]. These sources are listed in full in the "
                "References section."
            ),
        ],
    )

    replace_section(
        doc,
        "Figures and tables (if necessary)",
        "Details of proposed solution",
        [
            (
                "Fig. 4.1 illustrates the validation workflow: the rules defined on the data model "
                "are checked against the migrated records in the database, and the results are "
                "presented as a summary report and a pass or fail outcome for the release."
            ),
        ],
    )

    set_paragraph_text(
        doc.paragraphs[find_heading(doc, "Subsubsection if necessary")],
        "Sources and evidence base",
    )
    set_paragraph_text(
        doc.paragraphs[find_heading(doc, "Figures and tables (if necessary)")],
        "Figures and tables",
    )

    details_idx = find_heading(doc, "Details of proposed solution")
    anchor = doc.paragraphs[details_idx]
    for text, style in reversed(
        [
            (
                "Table 4.1. Validation markers available in v1 and their execution strategy",
                "Opis tabeli - Praca dyplomowa",
            ),
            (
                "Fig. 4.1. MigrationValidator workflow — declare rules on ORM columns, "
                "scan existing rows with SQL, report violations [8]",
                "Opis rysunku - Praca dyplomowa",
            ),
        ]
    ):
        new_p = deepcopy(anchor._p)
        anchor._p.addprevious(new_p)
        new_para = Paragraph(new_p, anchor._parent)
        new_para.style = style
        set_paragraph_text(new_para, text)

    fill_table(doc)

    replace_section(doc, "Details of proposed solution", "Summary and conclusions", [])

    append_section_content(
        doc,
        "Details of proposed solution",
        [
            (
                "The solution is a data-validation tool that runs immediately after a database "
                "migration and confirms that the migrated data still satisfies the business rules "
                "defined for it. Validation rules are attached to the data model itself, so they "
                "stay aligned with the application and require no separate, parallel configuration "
                "to maintain."
            ),
            (
                "This section sets out the problem the tool addresses, how the solution operates, the "
                "validation rules it supports, the reports it produces, how it integrates into the "
                "delivery process, a demonstration of it in use, the quality-assurance approach "
                "behind it, and the failure modes it covers today against those planned for future "
                "releases."
            ),
        ],
        [
            (
                "Problem addressed",
                [
                    "Standard migration checks confirm that data has been moved — matching record "
                    "counts, clean logs, successful connections — but they do not confirm that each "
                    "value still means what it did in the source system. As documented in the Sprint "
                    "I research, the most damaging migration failures pass every one of these checks: "
                    "the totals reconcile and the process reports success, yet values have been "
                    "altered, and the problem only surfaces weeks later through a customer or a "
                    "regulator.",
                    "A further gap is that most automatic safeguards only apply to new data as it is "
                    "entered through the application. Large migrations load data in bulk, outside "
                    "those safeguards, so the migrated records are never re-checked against the rules "
                    "the business expects them to follow.",
                    "The tool closes this gap by validating the data that is already in the target "
                    "database after a migration, at full production volume, against the rules defined "
                    "on the data model.",
                ],
            ),
            (
                "How the solution operates",
                [
                    "Validation rules are declared once, directly on the relevant fields of the data "
                    "model, and are reviewed and version-controlled together with the rest of the "
                    "application code. Because the rules live with the model, there is a single "
                    "source of truth and no risk of a separate rulebook drifting out of step with "
                    "the schema.",
                    "When the validation runs, the database performs the search for non-compliant "
                    "records and returns only those that breach a rule. This keeps the check fast and "
                    "allows it to operate on very large tables, since the tool processes results "
                    "efficiently rather than loading entire tables into memory.",
                    "The outcome is a structured list of every record that failed validation, "
                    "identifying the affected table, record, and field, together with the rule that "
                    "was breached.",
                ],
            ),
            (
                "Validation rules supported",
                [
                    "Mandatory field: the field must always contain a value and may never be left "
                    "empty.",
                    "Optional field: the field is explicitly permitted to be empty, so a blank value "
                    "is not reported as an error.",
                    "Unique value: no two records may share the same value in that field, such as a "
                    "customer identifier or email address.",
                    "Unique combination: a set of fields must be unique when taken together, for "
                    "cases where uniqueness depends on more than one field.",
                    "Minimum and maximum value: a numeric field must fall within an acceptable range, "
                    "for example a value that must not be negative or must not exceed a defined "
                    "ceiling.",
                    "Valid number: numeric fields must hold meaningful values, rejecting the "
                    "corrupted entries that can be introduced by data-processing tools.",
                ],
            ),
            (
                "Reporting and outputs",
                [
                    "The primary output is a clear table listing every record that failed validation, "
                    "with the table, record identifier, field, and a readable description of the "
                    "problem — suitable for handing directly to the team responsible for correction.",
                    "A summary view aggregates the failures by rule and field, showing how many "
                    "records breached each rule. This gives project and delivery managers an "
                    "at-a-glance picture of data quality without reviewing individual records.",
                    "The results can also be produced as a self-contained report with a clear "
                    "pass or fail status, providing a documented sign-off artefact to accompany a "
                    "migration before it is approved for production.",
                ],
            ),
            (
                "Integration into the delivery process",
                [
                    "The validation integrates with the automated checks that teams already run "
                    "before a release, where it acts as a single pass-or-fail quality gate on the "
                    "migrated data.",
                    "If any record breaches a rule, the release is blocked and the responsible team "
                    "is shown precisely what failed. Data-quality problems are therefore caught "
                    "before go-live rather than discovered in production, with no manual reporting "
                    "step to remember.",
                ],
            ),
            (
                "Demonstration",
                [
                    "To demonstrate the tool in a controlled setting, we created a sample database of "
                    "approximately 1,000 records and deliberately introduced the kinds of defects a "
                    "real bulk import can produce: 50 records with a missing email address, 20 "
                    "records with an invalid negative age, and 25 duplicated records.",
                    "Running the validation against this dataset identified each of these issues "
                    "correctly and grouped them by type, confirming that the tool detects exactly the "
                    "categories of error that conventional count-based migration checks overlook.",
                    "The summary produced by that run is shown below (rule, field, number of records "
                    "affected):",
                    "missing value     email     50",
                    "duplicate value   email     50",
                    "below minimum     age       20",
                    "The same demonstration is also available as an interactive notebook and as a "
                    "report-generating script, so the results can be reviewed in whichever format is "
                    "most useful to the audience.",
                ],
            ),
            (
                "Quality assurance",
                [
                    "The tool is supported by an automated test suite of more than thirty tests, "
                    "covering each validation rule individually as well as the complete process from "
                    "start to finish. Test coverage is measured automatically as part of the build.",
                    "Keeping the tool focused and thoroughly tested is a deliberate choice: it makes "
                    "the solution dependable and straightforward to maintain, which supports the "
                    "open-source, community-supported model set out in the business plan.",
                ],
            ),
            (
                "Failure modes covered and roadmap",
                [
                    "Covered today: missing values left behind by bulk imports; numeric values "
                    "outside an acceptable range; duplicated records from repeated or non-idempotent "
                    "jobs; broken unique combinations in systems serving many customers; and "
                    "corrupted numeric values introduced by data tooling. These map directly to the "
                    "highest-frequency structural failures identified in the Sprint I research.",
                    "Planned for future releases: detection of shifted dates and time-zone errors, "
                    "categories that silently change meaning, text corrupted by encoding changes, and "
                    "broken links between related tables. The solution was designed so these "
                    "additional checks can be added without re-engineering the existing tool.",
                ],
            ),
        ],
    )

    replace_section(
        doc,
        "Summary and conclusions",
        "Appendixes",
        [
            (
                "The project met the problem of undetected data corruption after successful database "
                "migrations. Research in Sprint I established motivation from real incidents; "
                "implementation in Sprint II produced a working SQLAlchemy validation library."
            ),
            (
                "Met problems: bulk loads bypass ORM checks; row counts do not prove correctness; "
                "teams discover defects via users rather than monitoring."
            ),
            (
                "Proposed solutions: Annotated markers on models; MigrationValidator SQL audit; "
                "pandas/HTML reporting; pytest CI gate."
            ),
            (
                "Unsolved issues: semantic validators (timezone, category remapping, encoding), "
                "cross-table referential integrity, sequence state after replication."
            ),
            (
                "Future work: domain-specific markers via parser.py; baseline diff against source; "
                "PostgreSQL/MySQL dialect optimisations; packaging for PyPI distribution."
            ),
            (
                "Team contributions: Hubert led engine and reporting; Patryk owned metadata/parser; "
                "Adrian implemented SQL validators; Dominik delivered pytest integration and demos; "
                "Beniamin produced research artefacts and documentation."
            ),
        ],
    )

    replace_section(
        doc,
        "Appendixes",
        "References",
        [
            "Appendix A — Repository layout: src/, Tests/, examples/, WriteUps/.",
            "Appendix B — Run instructions: uv run pytest; uv run python examples/validate_dirty_database.py.",
            "Appendix C — Presentation materials: WriteUps/why-and-when-slides.pdf and ice-sqlalchemy-validation-slides.pptx.",
        ],
    )

    ref_start = find_heading(doc, "References") + 1
    references = [
        "[1]\tNetflix Technology Blog, “Netflix Billing Migration to AWS — Part III,” 19 April 2017.",
        "[2]\tVovance, “Why Most Data Migration Projects Fail,” Medium, December 2025.",
        "[3]\tDev Engineer, “The PostgreSQL Migration That Corrupted Every Timestamp,” Medium, June 2025.",
        "[4]\tNewton, “What Broke the Bank,” Increment, 2019; FCA/PRA Final Notice, December 2022.",
        "[5]\tDan Luu, postmortems collection; OpenAI, “March 20 ChatGPT outage,” March 2023.",
        "[6]\tUK AAIB report on TUI Aviation load-sheet incident, 2019.",
        "[7]\tSQLAlchemy 2.0 documentation — ORM validation and mapped column typing.",
        "[8]\tICE SQLAlchemy Validation repository: src/, Tests/, examples/, README.md.",
        "[9]\tWriteUps/Error_refrences_SprintI/migration-data-quality-report.typ — Sprint I research report.",
        "[10]\tWriteUps/why-and-when-slides.pdf — companion presentation on library adoption.",
    ]
    ref_paras = [p for p in doc.paragraphs[ref_start:] if p.text.strip()]
    for i, text in enumerate(references):
        if i < len(ref_paras):
            set_paragraph_text(ref_paras[i], text)
        else:
            insert_after(ref_paras[-1], text, "Literatura - Praca Dyplomowa")
            ref_paras = [p for p in doc.paragraphs[ref_start:] if p.text.strip()]
    for p in ref_paras[len(references) :]:
        clear_paragraph(p)

    try:
        doc.save(OUTPUT)
        print(f"Saved: {OUTPUT}")
    except PermissionError:
        fallback = OUTPUT.with_name("ICE_projects_report_filled_new.docx")
        doc.save(fallback)
        print(f"{OUTPUT.name} was locked (open in Word?). Saved to: {fallback}")


if __name__ == "__main__":
    main()
